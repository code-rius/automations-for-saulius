from pathlib import Path
from dotenv import load_dotenv
from docx import Document
from docx.shared import Emu
from io import BytesIO
import zipfile
import xml.etree.ElementTree as ET
import copy
import os
import sys

load_dotenv()

DIR_SUTVARKYMAS = os.environ.get("DIR_SUTVARKYMAS")
DIR_PARASAS = os.environ.get("DIR_PARASAS")

if not DIR_SUTVARKYMAS:
    print("ENV missing: set DIR_SUTVARKYMAS in .env to the folder with .docx files")
    sys.exit(1)


def is_bullet_para(para):
    text = (para.text or "").strip()
    if not text:
        return False

    if text.startswith("•") or text.startswith("-"):
        return True

    try:
        style_name = (para.style.name or "").lower()
        if "bullet" in style_name or "list" in style_name:
            return True
    except Exception:
        pass

    try:
        pPr = para._p.pPr
        if pPr is not None and getattr(pPr, "numPr", None) is not None:
            return True
    except Exception:
        pass

    return False


def resolve_signature_path() -> Path | None:
    if not DIR_PARASAS:
        print("DIR_PARASAS not set, skipping signature.")
        return None

    p = Path(DIR_PARASAS)
    if p.is_dir():
        sig_docx = p / "Signature.docx"
        if sig_docx.exists():
            return sig_docx
        print(f"Signature.docx not found in {p}")
        return None
    else:
        if p.exists():
            return p
        print(f"Signature file not found: {p}")
        return None


def build_rel_map(sig_docx_path: Path):
    """Map rId -> internal media path (word/media/...) from signature docx."""
    rel_map = {}
    with zipfile.ZipFile(sig_docx_path, "r") as zf:
        rels_path = "word/_rels/document.xml.rels"
        if rels_path not in zf.namelist():
            return rel_map
        rels_xml = zf.read(rels_path)
    root = ET.fromstring(rels_xml)
    REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
    for rel in root.findall(f"{{{REL_NS}}}Relationship"):
        rId = rel.attrib.get("Id")
        target = rel.attrib.get("Target")
        if target and target.startswith("media/"):
            rel_map[rId] = "word/" + target
    return rel_map


def extract_image_bytes(sig_docx_path: Path, internal_path: str) -> bytes | None:
    with zipfile.ZipFile(sig_docx_path, "r") as zf:
        if internal_path in zf.namelist():
            return zf.read(internal_path)
    return None


def clear_cell_keep_props(cell):
    """Remove all content from a cell but keep tcPr (width, borders, etc.)."""
    tc = cell._tc
    tcPr = tc.tcPr
    for child in list(tc):
        if child is tcPr:
            continue
        tc.remove(child)
    cell.add_paragraph("")  # at least one paragraph


def copy_runs_with_images(dst_para, src_para, rel_map, sig_docx_path: Path):
    """
    Copy all runs from src_para to dst_para.
    If a run has an image, insert that image into dst_para with same size.
    """
    for run in src_para.runs:
        r_el = run._r

        blips = r_el.xpath(".//a:blip")
        if blips:
            rId = blips[0].get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            if rId and rId in rel_map:
                img_internal_path = rel_map[rId]
                img_bytes = extract_image_bytes(sig_docx_path, img_internal_path)
                if img_bytes:
                    extent_elems = r_el.xpath(".//wp:extent")
                    width_emu = None
                    height_emu = None
                    if extent_elems:
                        width_emu = int(extent_elems[0].get("cx"))
                        height_emu = int(extent_elems[0].get("cy"))
                    bio = BytesIO(img_bytes)
                    pic_run = dst_para.add_run()
                    if width_emu and height_emu:
                        pic_run.add_picture(bio, width=Emu(width_emu), height=Emu(height_emu))
                    else:
                        pic_run.add_picture(bio)
            continue

        # normal text run
        new_r = dst_para.add_run(run.text)
        new_r.bold = run.bold
        new_r.italic = run.italic
        new_r.underline = run.underline
        if run.font is not None:
            new_r.font.name = run.font.name
            new_r.font.size = run.font.size
            if run.font.color is not None:
                new_r.font.color.rgb = run.font.color.rgb


def insert_signature_table_preserving_layout(main_doc: Document, sig_docx_path: Path):
    """
    1. deep-copy the first table from signature doc into main doc (keeps layout)
    2. then replace cell content with text+images from the signature doc
    """
    sig_doc = Document(sig_docx_path)
    if not sig_doc.tables:
        # fallback: no table, just copy paragraphs
        rel_map = build_rel_map(sig_docx_path)
        for para in sig_doc.paragraphs:
            new_p = main_doc.add_paragraph()
            new_p.alignment = para.alignment
            copy_runs_with_images(new_p, para, rel_map, sig_docx_path)
        return True

    sig_table = sig_doc.tables[0]
    rel_map = build_rel_map(sig_docx_path)

    # 1) deep-copy original table XML so we keep layout
    tbl_xml = copy.deepcopy(sig_table._tbl)
    main_doc._body._element.append(tbl_xml)

    # 2) get the just-appended table
    new_table = main_doc.tables[-1]

    # 3) walk rows/cells by index (NOT sig_table.columns)
    for r_idx, sig_row in enumerate(sig_table.rows):
        new_row = new_table.rows[r_idx]
        for c_idx, _ in enumerate(sig_row.cells):
            src_cell = sig_table.cell(r_idx, c_idx)
            dst_cell = new_row.cells[c_idx]

            # clear content but keep cell properties
            clear_cell_keep_props(dst_cell)

            # copy paragraphs
            for para in src_cell.paragraphs:
                new_p = dst_cell.add_paragraph()
                new_p.alignment = para.alignment
                copy_runs_with_images(new_p, para, rel_map, sig_docx_path)

    return True


def add_spacing_paragraphs(doc: Document, count: int = 4):
    for _ in range(count):
        doc.add_paragraph()


def move_pridedama_and_clean(doc_path: Path, sig_path: Path | None):
    doc = Document(doc_path)
    changed = False

    # 1) move 'Pridedama' above bullets
    paras = list(doc.paragraphs)
    pr_indices = [i for i, p in enumerate(paras) if "pridedama" in (p.text or "").lower()]

    for pr_idx in reversed(pr_indices):
        paras = list(doc.paragraphs)
        pr_para = paras[pr_idx]

        start = pr_idx - 1
        if start < 0:
            continue

        while start >= 0 and is_bullet_para(paras[start]):
            start -= 1

        first_bullet_idx = start + 1

        if first_bullet_idx <= pr_idx - 1:
            body = doc._body._element
            pr_elem = pr_para._p
            parent = pr_elem.getparent()
            if parent is None:
                continue
            parent.remove(pr_elem)

            bullet_elem = paras[first_bullet_idx]._p
            insert_pos = body.index(bullet_elem)
            body.insert(insert_pos, pr_elem)
            changed = True
            print(f"Moved 'Pridedama:' in {doc_path.name}")

    # 2) remove everything after last bullet
    paras = list(doc.paragraphs)
    last_bullet_idx = -1
    for i, p in enumerate(paras):
        if is_bullet_para(p):
            last_bullet_idx = i

    if last_bullet_idx != -1 and last_bullet_idx < len(paras) - 1:
        body = doc._body._element
        for p in paras[last_bullet_idx + 1:]:
            elem = p._p
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)
                changed = True
        print(f"Removed text after last bullet in: {doc_path.name}")

    # 3) add spacing
    add_spacing_paragraphs(doc, 3)
    changed = True

    # 4) insert signature table with preserved layout
    if sig_path is not None:
        ok = insert_signature_table_preserving_layout(doc, sig_path)
        if ok:
            changed = True
            print(f"Signature (table) added to: {doc_path.name}")
    else:
        print(f"No signature added to: {doc_path.name}")

    if changed:
        doc.save(doc_path)
        print(f"Fixed: {doc_path.name}")
    else:
        print(f"No change: {doc_path.name}")


def main():
    folder = Path(DIR_SUTVARKYMAS)
    if not folder.exists() or not folder.is_dir():
        print(f"Path not found or not a directory: {folder}")
        return

    sig_path = resolve_signature_path()

    docx_files = list(folder.glob("*.docx"))
    if not docx_files:
        print("No .docx files found in folder.")
        return

    for f in docx_files:
        try:
            move_pridedama_and_clean(f, sig_path)
        except Exception as e:
            print(f"Error processing {f.name}: {e}")


if __name__ == "__main__":
    main()
