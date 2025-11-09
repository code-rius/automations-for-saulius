from pathlib import Path
from dotenv import load_dotenv
from docx import Document
import os
import copy
import sys

load_dotenv()

DIR_SUTVARKYMAS = os.environ.get("DIR_SUTVARKYMAS")
if not DIR_SUTVARKYMAS:
    print("ENV missing: set DIR_SUTVARKYMAS in .env to the folder with .docx files")
    sys.exit(1)

def is_bullet_para(para):
    """Heuristics to detect a bullet/list paragraph."""
    text = (para.text or "").strip()
    if not text:
        return False
    # common visible bullet char
    if text.startswith("•") or text.startswith("-"):
        return True
    # style name often contains 'bullet' or 'list'
    try:
        style_name = (para.style.name or "").lower()
        if "bullet" in style_name or "list" in style_name:
            return True
    except Exception:
        pass
    # try to detect numbering in the underlying XML (numPr present)
    try:
        pPr = para._p.pPr  # may raise AttributeError
        if pPr is not None and getattr(pPr, "numPr", None) is not None:
            return True
    except Exception:
        pass
    return False

def move_pridedama_before_bullets(doc_path: Path):
    doc = Document(doc_path)
    paras = list(doc.paragraphs)
    changed = False

    # find all paragraphs that contain 'Pridedama:' (case-insensitive)
    pr_indices = [i for i, p in enumerate(paras) if "pridedama:" in (p.text or "").lower()]

    # process from last to first to avoid index shift issues
    for pr_idx in reversed(pr_indices):
        # find contiguous bullet block immediately before pr_idx
        start = pr_idx - 1
        if start < 0:
            continue
        # move backwards while paragraphs are bullets (stop at first non-bullet)
        while start >= 0 and is_bullet_para(paras[start]):
            start -= 1
        first_bullet_idx = start + 1

        # if there is at least one bullet before Pridedama, move Pridedama to before first_bullet_idx
        if first_bullet_idx <= pr_idx - 1:
            pr_para = paras[pr_idx]
            # copy element and insert before the first bullet
            pr_elem = pr_para._p
            pr_copy = copy.deepcopy(pr_elem)
            doc._body._element.insert(first_bullet_idx, pr_copy)
            # remove the original pr element (it may have shifted; remove by reference)
            parent = pr_elem.getparent()
            if parent is not None:
                parent.remove(pr_elem)
            changed = True
            # refresh paras list for subsequent iterations
            paras = list(doc.paragraphs)

    if changed:
        doc.save(doc_path)
        print(f"Fixed: {doc_path.name}")
    else:
        print(f"No change: {doc_path.name}")

def main():
    folder = Path(DIR_SUTVARKYMAS)
    if not folder.exists() or not folder.is_dir():
        print(f"Path not found or not a directory: {folder}")
        return

    docx_files = list(folder.glob("*.docx"))
    if not docx_files:
        print("No .docx files found in folder.")
        return

    for f in docx_files:
        try:
            move_pridedama_before_bullets(f)
        except Exception as e:
            print(f"Error processing {f.name}: {e}")

if __name__ == "__main__":
    main()