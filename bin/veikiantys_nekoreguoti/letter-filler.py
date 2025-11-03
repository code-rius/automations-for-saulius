"""
Letter Generator Script

This script generates customized letters based on a template document and CSV data.
It handles proper formatting, signatures, and project-specific content.
"""

from pathlib import Path
import csv
import os
import copy
from datetime import date
from collections import defaultdict
import re

from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

load_dotenv()

class FormatHelper:
    """Helper class for document text and formatting operations"""
    
    @staticmethod
    def replace_text_in_paragraph(paragraph, replacements):
        """Replace placeholders in a paragraph while preserving formatting."""
        for key, value in replacements.items():
            for run in paragraph.runs:
                if key in run.text:
                    run.text = run.text.replace(key, value)
    
    @staticmethod
    def replace_text_in_cell(cell, replacements):
        """Replace placeholders in a table cell while preserving formatting."""
        for paragraph in cell.paragraphs:
            FormatHelper.replace_text_in_paragraph(paragraph, replacements)
    
    @staticmethod
    def copy_run_format(source_run, target_run):
        """Copy formatting from one run to another."""
        target_run.bold = source_run.bold
        target_run.italic = source_run.italic
        target_run.underline = source_run.underline
        if source_run.font.size:
            target_run.font.size = source_run.font.size
        if source_run.font.name:
            target_run.font.name = source_run.font.name
        if hasattr(source_run.font, 'color') and source_run.font.color:
            target_run.font.color.rgb = source_run.font.color.rgb
    
    @staticmethod
    def copy_para_format(source_para, target_para):
        """Copy paragraph formatting attributes from one paragraph to another."""
        for attr in ['alignment', 'left_indent', 'right_indent', 'first_line_indent', 
                    'line_spacing', 'space_before', 'space_after', 'keep_together', 
                    'keep_with_next', 'page_break_before', 'widow_control']:
            if hasattr(source_para.paragraph_format, attr):
                setattr(target_para.paragraph_format, attr, getattr(source_para.paragraph_format, attr))
        
        # Copy style if available
        if source_para.style:
            target_para.style = source_para.style


class DocumentHelper:
    """Helper class for document manipulation operations"""
    
    @staticmethod
    def find_paragraph_with_text(doc, text_to_find):
        """Find a paragraph containing specific text."""
        for i, para in enumerate(doc.paragraphs):
            if text_to_find in para.text:
                return i
        return -1
    
    @staticmethod
    def set_bullet_numbering(para_with_numbering, target_para):
        """Copy bullet points/numbering from one paragraph to another."""
        if not hasattr(para_with_numbering, '_p') or para_with_numbering._p.pPr is None:
            return False
        
        p_pr = para_with_numbering._p.pPr
        if p_pr.numPr is None:
            return False
        
        # Get the numId and ilvl from source paragraph
        if p_pr.numPr.numId is None or p_pr.numPr.ilvl is None:
            return False
            
        num_id = p_pr.numPr.numId.val
        ilvl = p_pr.numPr.ilvl.val
        
        # Make sure target paragraph has a paragraph properties element
        if target_para._p.pPr is None:
            target_para._p.get_or_add_pPr()
        
        # Add numPr element if it doesn't exist
        num_pr = target_para._p.pPr.get_or_add_numPr()
        
        # Set the numId - identifies the numbering definition
        num_id_element = OxmlElement('w:numId')
        num_id_element.set(qn('w:val'), str(num_id))
        num_pr.append(num_id_element)
        
        # Set the ilvl - identifies the numbering level
        ilvl_element = OxmlElement('w:ilvl')
        ilvl_element.set(qn('w:val'), str(ilvl))
        num_pr.append(ilvl_element)
        
        return True


class LetterGenerator:
    """Class responsible for letter generation and content management"""
    
    def __init__(self, template_path):
        """Initialize with template document."""
        self.template_path = template_path
        self.template_doc = Document(template_path)
    
    def create_letter(self, recipient_data, plot_data, project_data, today_date):
        """Create a customized letter based on provided data."""
        # Create a new document from the template
        doc = Document(self.template_path)
        
        # Set up replacements dictionary
        replacements = {
            "gavejas_1": recipient_data["name"],
            "adresas_2": recipient_data["address"],
            "pasto_kodas_3": recipient_data["postal_code"],
            "proj_data": today_date
        }
        
        # Apply replacements to document text
        self._apply_replacements(doc, replacements)
        
        # Fill table with plot data
        self._fill_table_with_plots(doc, plot_data)
        
        # Add project descriptions
        self._add_project_descriptions(doc, project_data)
        
        # Add attestation paragraphs
        self._add_attestation_paragraphs(doc, project_data)
        
        # Copy signature and add email
        self._add_signature_content(doc)
        self._ensure_email_in_document(doc, "domantas.aleknavicius@etprojektai.eu")
        
        return doc
        
    def _apply_replacements(self, doc, replacements):
        """Apply text replacements throughout the document."""
        # Replace in paragraphs
        for para in doc.paragraphs:
            FormatHelper.replace_text_in_paragraph(para, replacements)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    FormatHelper.replace_text_in_cell(cell, replacements)
    
    def _fill_table_with_plots(self, doc, plots_data):
        """Fill the table with plot data, adding rows as needed."""
        if not doc.tables or not plots_data:
            return
        
        table = doc.tables[0]
        if len(table.rows) < 2:
            return
        
        # Get the template row (second row)
        template_row = table.rows[1]
        
        # Clear template row
        for cell in template_row.cells:
            cell.text = ""
        
        # Fill in first plot data
        if plots_data:
            first_plot = plots_data[0]
            template_row.cells[0].text = first_plot[0]  # Registro Nr
            template_row.cells[1].text = first_plot[2]  # Unikalus Nr
            template_row.cells[2].text = first_plot[3]  # Kadastro Nr
            template_row.cells[3].text = first_plot[1]  # Sklypo adresas
        
        # Add additional rows for the rest of the plots
        for plot in plots_data[1:]:
            new_row = table.add_row()
            new_row.cells[0].text = plot[0]  # Registro Nr
            new_row.cells[1].text = plot[2]  # Unikalus Nr
            new_row.cells[2].text = plot[3]  # Kadastro Nr
            new_row.cells[3].text = plot[1]  # Sklypo adresas
    
    def _add_project_descriptions(self, doc, project_data):
        """Add project descriptions to the document."""
        if not project_data:
            return
            
        # Find paragraphs containing key placeholders
        proj_pav_para_index = DocumentHelper.find_paragraph_with_text(doc, "proj_pav_5")
        elektrine_para_index = DocumentHelper.find_paragraph_with_text(doc, "elektrines_numeris_11")
        
        if proj_pav_para_index < 0 or elektrine_para_index < 0:
            return
            
        proj_pav_para = doc.paragraphs[proj_pav_para_index]
        elektrine_para = doc.paragraphs[elektrine_para_index]
        
        # Keep original project order for proj_pav_5 behavior
        project_items = list(project_data.items())
        first_elektrine_nr, first_project_info = project_items[0]
        
        # Replace proj_pav_5 with first project's name (unchanged)
        for run in proj_pav_para.runs:
            if "proj_pav_5" in run.text:
                run.text = run.text.replace("proj_pav_5", first_project_info.get("projekt_pav", ""))
        
        # Build sorted list of elektrine keys by numeric part after "VE"
        def _ve_sort_key(key):
            m = re.search(r"VE(\d+)", key, re.IGNORECASE)
            return int(m.group(1)) if m else float("inf")
        
        sorted_keys = sorted(project_data.keys(), key=_ve_sort_key)
        
        # Replace elektrines_numeris_11 placeholder with ascending-ordered generator list (comma separated)
        formatted_list = ", ".join(sorted_keys)
        for run in elektrine_para.runs:
            if "elektrines_numeris_11" in run.text:
                run.text = run.text.replace("elektrines_numeris_11", formatted_list)
        
        # Skip the first project as we've already processed it (proj_pav paragraphs preserved in template order)
        if len(project_items) <= 1:
            return
            
        # Process additional projects (append proj_pav paragraphs in original project_data order)
        index_for_proj = proj_pav_para_index + 1
        
        for elektrine_nr, project_info in project_items[1:]:
            # Create a new project paragraph (pass project name to include address)
            new_para = self._create_project_paragraph(doc, proj_pav_para, elektrine_nr, project_info.get("projekt_pav", ""))
            
            # Insert at the correct position and remove from end
            doc._body._element.insert(index_for_proj, new_para._element)
            doc._body._element.remove(doc.paragraphs[-1]._element)
            
            # Update index for next insertion
            index_for_proj += 1
    
    def _create_project_paragraph(self, doc, template_para, elektrine_nr, project_pav=""):
        """Create a formatted project paragraph including project name/address."""
        new_para = doc.add_paragraph()
        FormatHelper.copy_para_format(template_para, new_para)
        
        # Add opening quote
        quote_open = new_para.add_run("„")
        quote_open.font.name = "Arial"
        quote_open.font.size = Pt(11)
        
        # Normalize and compose content without duplicating full phrase
        prefix = "Energijos iš atsinaujinančių išteklių gamybos paskirties inžinerinio statinio, vėjo elektrinės"
        proj = (project_pav or "").strip()
        
        # If projekt_pav already contains the full prefix (possibly with elektrine_nr), use it (avoid duplication)
        if proj:
            if prefix.lower() in proj.lower():
                # If elektrine_nr missing inside projekt_pav, ensure it's present once
                if elektrine_nr and elektrine_nr not in proj:
                    # remove any leading/trailing punctuation to join cleanly
                    proj_clean = proj.strip(' ,;.')
                    content = f"{prefix} {elektrine_nr}, {proj_clean}"
                else:
                    content = proj
            else:
                # projekt_pav likely just location/address -> append into template phrase
                cleaned = re.sub(r"\s*[,;:]?\s*statybos projektas\.?$", "", proj, flags=re.IGNORECASE).strip(' ,;.')
                content = f"{prefix} {elektrine_nr}, {cleaned}, statybos projektas"
        else:
            # Fallback
            content = f"{prefix} {elektrine_nr}, statybos projektas"
        
        content_run = new_para.add_run(content)
        content_run.font.name = "Arial"
        content_run.font.size = Pt(11)
        
        # Add closing quote and semicolon/quote char matching template
        quote_close = new_para.add_run("\";")
        quote_close.font.name = "Arial"
        quote_close.font.size = Pt(11)
        
        return new_para
    
    def _add_attestation_paragraphs(self, doc, project_data):
        """Insert attestation bullets after 'Pridedama:' in the same order as project paragraphs.
        Use the template's attestation paragraph as source so each bullet gets a single VE number.
        """
        pridedama_index = DocumentHelper.find_paragraph_with_text(doc, "Pridedama:")
        if pridedama_index < 0:
            return

        # Determine VE order by scanning project paragraphs before Pridedama
        proj_prefix = "Energijos iš atsinaujinančių išteklių gamybos paskirties inžinerinio statinio"
        ordered_ves = []
        for i, para in enumerate(doc.paragraphs):
            if i >= pridedama_index:
                break
            text = para.text or ""
            if proj_prefix in text:
                m = re.findall(r"\bVE[0-9A-Za-z._-]*", text)
                if m:
                    ordered_ves.append(m[0])
        if not ordered_ves:
            ordered_ves = list(project_data.keys())

        # Remove existing attestation paragraphs immediately after Pridedama:
        start = pridedama_index + 1
        end = start
        while end < len(doc.paragraphs) and ("Skelbimas apie" in doc.paragraphs[end].text or "projektinių pasiūlymų viešinimą" in doc.paragraphs[end].text):
            end += 1
        for idx in range(end - 1, start - 1, -1):
            p = doc.paragraphs[idx]._p
            parent = p.getparent()
            if parent is not None:
                parent.remove(p)

        # Use original template attestation paragraph (keeps formatting clean)
        template_attestation_para = None
        for para in self.template_doc.paragraphs:
            if "Skelbimas apie" in para.text:
                template_attestation_para = para
                break
        if template_attestation_para is None:
            # fallback to using the current doc paragraph at pridedama_index (if any)
            if pridedama_index + 1 < len(doc.paragraphs):
                template_attestation_para = doc.paragraphs[pridedama_index + 1]
            else:
                template_attestation_para = None

        # Insert attestation bullets one VE per paragraph (in derived order)
        insert_pos = pridedama_index + 1
        for ve in ordered_ves:
            # create paragraph (appends one to doc) using template for formatting
            source_para = template_attestation_para if template_attestation_para is not None else doc.paragraphs[pridedama_index]
            new_att = self._create_attestation_paragraph(doc, source_para, ve)
            # insert deep copy at target position and remove appended original
            new_elem = new_att._p
            new_copy = copy.deepcopy(new_elem)
            doc._body._element.insert(insert_pos, new_copy)
            parent = new_elem.getparent()
            if parent is not None:
                parent.remove(new_elem)
            insert_pos += 1
    
    def _create_attestation_paragraph(self, doc, template_para, elektrine_nr):
        """Create a formatted attestation paragraph."""
        new_para = doc.add_paragraph()
        FormatHelper.copy_para_format(template_para, new_para)
        
        # Copy bullet points/numbering
        DocumentHelper.set_bullet_numbering(template_para, new_para)
        
        # Add content with Arial 11pt
        text = f"Skelbimas apie energijos iš atsinaujinančių išteklių gamybos paskirties inžinerinio statinio, vėjo elektrinės {elektrine_nr}, projektinių pasiūlymų viešinimą (2 lapai);"
        run = new_para.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        
        return new_para
    
    def _add_signature_content(self, doc):
        """Copy the signature content from template to the document."""
        # Find the "Pagarbiai," paragraph in template
        pagarbiai_idx_template = DocumentHelper.find_paragraph_with_text(self.template_doc, "Pagarbiai")
        if pagarbiai_idx_template < 0:
            return False
        
        # Find the last attestation paragraph in target document
        last_attestation_idx = -1
        pridedama_idx = DocumentHelper.find_paragraph_with_text(doc, "Pridedama:")
        
        if pridedama_idx >= 0:
            for i, para in enumerate(doc.paragraphs):
                if "Skelbimas apie" in para.text and i > pridedama_idx:
                    last_attestation_idx = i
        
        # Add spacing after attestations
        if last_attestation_idx >= 0:
            doc.add_paragraph()
            doc.add_paragraph()
        
        # Get signature paragraphs from template
        signature_paras = self.template_doc.paragraphs[pagarbiai_idx_template:]
        
        # Add signature paragraphs with exact formatting
        for para in signature_paras:
            new_para = doc.add_paragraph()
            FormatHelper.copy_para_format(para, new_para)
            
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                FormatHelper.copy_run_format(run, new_run)
        
        # Copy drawings from template to target document
        self._copy_drawing_objects(doc)
        
        return True
    
    def _copy_drawing_objects(self, doc):
        """Copy drawing objects from template to target document."""
        try:
            # Access template document part to find drawings
            template_doc_part = self.template_doc._part
            
            if not hasattr(template_doc_part, '_element'):
                return
                
            # Find all drawing objects
            drawings = template_doc_part._element.xpath('.//w:drawing')
            if not drawings:
                return
                
            # Find the "Pagarbiai," paragraph in target document
            pagarbiai_idx_target = DocumentHelper.find_paragraph_with_text(doc, "Pagarbiai")
            if pagarbiai_idx_target < 0:
                return
                
            # Get the paragraph element
            target_para = doc.paragraphs[pagarbiai_idx_target]._p
            
            # Clone and insert each drawing
            for drawing in drawings:
                drawing_copy = copy.deepcopy(drawing)
                target_para.append(drawing_copy)
                
        except Exception as e:
            print(f"Error copying drawing objects: {e}")
    
    def _ensure_email_in_document(self, doc, email_address="domantas.aleknavicius@etprojektai.eu"):
        """Ensure the email address is properly set in the document."""
        for i, para in enumerate(doc.paragraphs):
            if "El. p.:" in para.text:
                # Clear and rebuild the paragraph
                for run in list(para.runs):
                    run.clear()
                
                # Add label and email with Arial 11pt
                label_run = para.add_run("El. p.:")
                label_run.font.name = "Arial"
                label_run.font.size = Pt(11)
                
                email_run = para.add_run(" " + email_address)
                email_run.font.name = "Arial"
                email_run.font.size = Pt(11)
                
                return True
        
        return False


class CsvProcessor:
    """Class for processing CSV data into structured information for letters."""
    
    def __init__(self, csv_path):
        """Initialize with CSV file path."""
        self.csv_path = csv_path
        
    def read_data(self):
        """Read and process CSV data into individuals with their plots and projects."""
        # Detect delimiter
        with open(self.csv_path, "r", encoding="utf-8-sig") as f:
            sample = f.read(4096)
            dialect = csv.Sniffer().sniff(sample)
            delimiter = dialect.delimiter
        
        # Read CSV data
        with open(self.csv_path, "r", encoding="utf-8-sig") as f:
            reader = csv.reader(f, delimiter=delimiter)
            header = next(reader)  # Skip header row
            rows = list(reader)
            
        # Group by individual
        individuals = defaultdict(list)
        
        for row in rows:
            if len(row) < 9:  # Need at least through the Tipas column
                continue
                
            tipas = row[8].lower()
            
            if tipas == "fizinis":
                # For individuals: combine first name and last name
                vardas = row[5]
                pavarde = row[6]
                id_or_date = row[7]
                individual_key = (vardas, pavarde, id_or_date)
            elif tipas == "juridinis":
                # For companies: use only the company name
                vardas = row[5]
                pavarde = ""
                id_or_date = row[7]
                individual_key = (vardas, "", id_or_date)
            else:
                continue
                
            individuals[individual_key].append(row)
            
        return individuals
        
    def process_individual(self, individual_key, individual_rows):
        """Process data for a single individual."""
        vardas, pavarde, id_or_date = individual_key
        tipas = individual_rows[0][8].lower()
        
        # Skip if all entries have no address
        if all(len(row) <= 12 or not row[12] for row in individual_rows):
            return None
            
        # Get the first row with an address
        address_row = next((row for row in individual_rows if len(row) > 12 and row[12]), None)
        if not address_row:
            return None
            
        # Create recipient data
        recipient_data = {
            "name": f"{vardas} {pavarde}" if tipas == "fizinis" else vardas,
            "address": address_row[12],
            "postal_code": address_row[13] if len(address_row) > 13 and address_row[13] else ""
        }
        
        # Collect unique projects and plot data
        projects = {}
        plots = set()
        
        for row in individual_rows:
            # Clean up the elektrine_nr 
            elektrine_nr = row[9].strip().replace('\ufeff', '')
            
            # Only add unique projects
            if elektrine_nr not in projects:
                projects[elektrine_nr] = {
                    "projekt_nr": row[10],
                    "projekt_pav": row[11]
                }
            
            # Add unique plots
            plot_tuple = (row[0], row[1], row[2], row[3])
            plots.add(plot_tuple)
            
        return {
            "recipient": recipient_data,
            "projects": projects,
            "plots": list(plots)
        }


def main():
    # Load environment variables
    etapas_dir = os.environ.get("DIR_ETAPAS")
    template_filename = os.environ.get("TEMPLATE_FILE_NAME")
    csv_filename = os.environ.get("ETAPAS_OUTPUT_FILE_NAME", "aggregated_output.csv")
    
    if not etapas_dir or not template_filename:
        print("Please set DIR_ETAPAS and TEMPLATE_FILE_NAME in your .env file.")
        exit(1)
    
    etapas_path = Path(etapas_dir)
    template_path = etapas_path / template_filename
    csv_path = etapas_path / csv_filename
    
    # Check if files exist
    if not template_path.exists() or not csv_path.exists():
        print(f"Error: Required files not found.")
        print(f"Template: {template_path} - {'Exists' if template_path.exists() else 'Missing'}")
        print(f"CSV: {csv_path} - {'Exists' if csv_path.exists() else 'Missing'}")
        exit(1)
    
    print(f"Using template: {template_path}")
    print(f"Reading data from: {csv_path}")
    
    # Create output folder
    output_folder = etapas_path / "letters"
    output_folder.mkdir(exist_ok=True)
    
    # Get today's date
    today_date = date.today().strftime("%Y-%m-%d")
    
    # Initialize processors
    csv_processor = CsvProcessor(csv_path)
    letter_generator = LetterGenerator(template_path)
    
    # Read individual data
    individuals = csv_processor.read_data()
    print(f"Found {len(individuals)} unique individuals/companies")
    
    # Process each individual
    processed_count = 0
    
    for individual_key, individual_rows in individuals.items():
        vardas, pavarde, _ = individual_key
        
        # Process data for this individual
        data = csv_processor.process_individual(individual_key, individual_rows)
        if not data:
            print(f"Skipping {vardas} {pavarde} - insufficient data")
            continue
            
        # Create document
        doc = letter_generator.create_letter(
            data["recipient"],
            data["plots"],
            data["projects"],
            today_date
        )
        
        # Generate filename and save
        safe_name = data["recipient"]["name"].replace(" ", "_").replace("/", "-").replace('"', '')
        output_filename = f"{safe_name}.docx"
        output_path = output_folder / output_filename
        
        doc.save(output_path)
        processed_count += 1
        
        print(f"Created document: {output_filename} (with {len(data['projects'])} projects and {len(data['plots'])} plots)")
    
    print(f"\nGenerated {processed_count} documents in: {output_folder}")


if __name__ == "__main__":
    main()