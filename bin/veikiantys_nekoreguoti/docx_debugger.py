from pathlib import Path
from docx import Document
from dotenv import load_dotenv
import os
import sys

load_dotenv()

def print_doc_structure(doc_path):
    """Print the structure and content of a .docx file to the terminal."""
    print(f"{'='*50}")
    print(f"DOCUMENT: {doc_path}")
    print(f"{'='*50}")
    
    try:
        doc = Document(doc_path)
    except Exception as e:
        print(f"ERROR opening document: {e}")
        return
    
    # Print document paragraphs
    print("\nPARAGRAPHS:")
    print(f"{'-'*50}")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():  # Only print non-empty paragraphs
            print(f"Para {i+1}: {para.text}")
            
            # Show run information for debugging formatting
            if len(para.runs) > 1:
                print(f"  - Contains {len(para.runs)} runs:")
                for j, run in enumerate(para.runs):
                    print(f"    Run {j+1}: '{run.text}' [Bold: {run.bold}, Italic: {run.italic}]")
    
    # Print document tables
    print("\nTABLES:")
    print(f"{'-'*50}")
    for i, table in enumerate(doc.tables):
        print(f"Table {i+1}: {len(table.rows)} rows x {len(table.columns)} columns")
        
        # Print table contents
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                text = cell.text.replace('\n', ' ')
                print(f"  Cell [{row_idx+1},{col_idx+1}]: {text}")
    
    # Print information about inline shapes (including images)
    print("\nINLINE SHAPES:")
    print(f"{'-'*50}")
    try:
        if hasattr(doc, 'inline_shapes'):
            shapes = doc.inline_shapes
            if shapes:
                for i, shape in enumerate(shapes):
                    shape_type = "Unknown"
                    if shape.type == 1:
                        shape_type = "Linked Picture"
                    elif shape.type == 2:
                        shape_type = "Embedded Picture"
                    elif shape.type == 3:
                        shape_type = "Chart"
                    elif shape.type == 4:
                        shape_type = "Smart Art"
                    elif shape.type == 5:
                        shape_type = "Embedded Object"
                    
                    width = shape.width if hasattr(shape, 'width') else "N/A"
                    height = shape.height if hasattr(shape, 'height') else "N/A"
                    
                    print(f"Shape {i+1}: Type={shape_type}, Width={width}, Height={height}")
            else:
                print("No inline shapes found")
    except Exception as e:
        print(f"Error checking inline shapes: {e}")
    
    # Check for other types of shapes (Drawing objects)
    print("\nDRAWING OBJECTS:")
    print(f"{'-'*50}")
    try:
        # Access the document part to find drawings
        main_document_part = doc._part
        if hasattr(main_document_part, '_element'):
            # Find all drawing objects using XPath
            drawings = main_document_part._element.xpath('.//w:drawing')
            if drawings:
                print(f"Found {len(drawings)} drawing objects in document")
                for i, drawing in enumerate(drawings):
                    # Try to identify which paragraph contains this drawing
                    para_index = "Unknown"
                    for j, para in enumerate(doc.paragraphs):
                        if para._p.xpath('.//w:drawing'):
                            for draw in para._p.xpath('.//w:drawing'):
                                if draw == drawing:
                                    para_index = j + 1
                                    break
                    
                    print(f"Drawing {i+1}: Located in paragraph approximately #{para_index}")
            else:
                print("No drawing objects found")
    except Exception as e:
        print(f"Error checking drawing objects: {e}")

def main():
    # Priority 1: Command line argument
    if len(sys.argv) > 1:
        doc_path = sys.argv[1]
        print(f"Using document path from command line: {doc_path}")
    
    # Priority 2: DEBUG_DOXC_PATH from .env
    else:
        debug_path = os.environ.get("DEBUG_DOXC_PATH")
        if debug_path:
            doc_path = debug_path
            print(f"Using document path from DEBUG_DOXC_PATH: {doc_path}")
        
        # Priority 3: Template path from .env
        else:
            etapas_dir = os.environ.get("DIR_ETAPAS")
            template_filename = os.environ.get("TEMPLATE_FILE_NAME")
            
            if not etapas_dir or not template_filename:
                print("No document path provided. Please either:")
                print("1. Provide a path as command line argument")
                print("2. Set DEBUG_DOXC_PATH in .env")
                print("3. Set DIR_ETAPAS and TEMPLATE_FILE_NAME in .env")
                return
                
            doc_path = Path(etapas_dir) / template_filename
            print(f"Using template path: {doc_path}")
    
    print_doc_structure(doc_path)

if __name__ == "__main__":
    main()