from pathlib import Path
from docx import Document
from docx.shared import Pt

docx_path = Path(__file__).parent.parent / "resources" / "sablonas.docx"
doc = Document(docx_path)
table = doc.tables[0]

# Reference row for formatting
template_row = table.rows[0]

# Add a new row
new_row = table.add_row()

# Copy text and formatting from the template row
for i, cell in enumerate(template_row.cells):
    new_cell = new_row.cells[i]
    
    # Copy paragraph text and formatting
    new_cell.text = cell.text
    for j, paragraph in enumerate(cell.paragraphs):
        new_para = new_cell.paragraphs[j]
        for k, run in enumerate(paragraph.runs):
            new_run = new_para.runs[k]
            new_run.font.name = run.font.name
            new_run.font.size = run.font.size
            new_run.font.bold = run.font.bold
            new_run.font.italic = run.font.italic
            new_run.font.color.rgb = run.font.color.rgb

# Save
doc.save("modified_with_row.docx")

for para in doc.paragraphs:
    print(para.text)

for table in doc.tables:
    for row in table.rows:
        row_text = [cell.text for cell in row.cells]
        print(", ".join(row_text))